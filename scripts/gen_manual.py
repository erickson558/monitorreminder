"""Generate the MonitorReminder user manual as a Word document."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(BASE, "docs", "MonitorReminder_Manual_Usuario.docx")

# ─── colour constants ────────────────────────────────────────────────────────
NAVY   = RGBColor(0x08, 0x1b, 0x2d)   # dark navy – headings
CYAN   = RGBColor(0x1a, 0x86, 0xb8)   # teal – accent
TEAL   = RGBColor(0x24, 0x70, 0x6c)   # green-teal – tips
GRAY   = RGBColor(0x44, 0x44, 0x44)   # body text
RED    = RGBColor(0xb0, 0x3d, 0x3d)   # warnings

# ─── helpers ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = NAVY
    p.runs[0].font.size = Pt(20)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = CYAN
    p.runs[0].font.size = Pt(14)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = NAVY
    p.runs[0].font.size = Pt(12)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    return p


def body(doc, text, bold=False, color=GRAY):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = color
    if bold:
        r.bold = True
    p.paragraph_format.space_after = Pt(4)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = GRAY
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    return p


def numbered(doc, text, level=0):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = GRAY
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    return p


def tip(doc, text):
    """Green callout box for tips."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("💡  " + text)
    r.font.size = Pt(11)
    r.font.color.rgb = TEAL
    r.italic = True
    return p


def warn(doc, text):
    """Red callout for warnings/limitations."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("⚠️  " + text)
    r.font.size = Pt(11)
    r.font.color.rgb = RED
    r.bold = True
    return p


def kbd(doc, text):
    """Inline keyboard shortcut style."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = "Consolas"
    r.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(2)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        set_cell_bg(cell, "1a86b8")

    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()
    return table


def divider(doc):
    p = doc.add_paragraph("─" * 80)
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


# ─── document ────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# Default font
doc.styles["Normal"].font.name = "Segoe UI"
doc.styles["Normal"].font.size = Pt(11)
doc.styles["Normal"].font.color.rgb = GRAY

# ── COVER ─────────────────────────────────────────────────────────────────────
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover.add_run("MonitorReminder")
r.font.size = Pt(36)
r.font.bold = True
r.font.color.rgb = NAVY

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Manual de Usuario")
r2.font.size = Pt(20)
r2.font.color.rgb = CYAN

doc.add_paragraph()
ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = ver.add_run("Versión 0.3.2  ·  2026")
r3.font.size = Pt(12)
r3.font.color.rgb = GRAY
r3.italic = True

doc.add_paragraph()
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = sub2.add_run("Centro orbital para sincronizar perfiles de ventanas por monitor")
r4.font.size = Pt(12)
r4.font.color.rgb = TEAL
r4.italic = True

doc.add_page_break()

# ── 1. INTRODUCCIÓN ───────────────────────────────────────────────────────────
h1(doc, "1. Introducción")
body(doc,
    "MonitorReminder es una utilidad de escritorio para Windows que recuerda dónde "
    "tenías abiertas tus ventanas y las devuelve exactamente a esa posición cuando "
    "cambias de configuración de monitores, reconectas una pantalla externa, o "
    "simplemente quieres restablecer tu espacio de trabajo."
)
body(doc,
    "Con MonitorReminder puedes:"
)
bullet(doc, "Guardar hasta 5 perfiles de distribución de ventanas.")
bullet(doc, "Restaurar cualquier perfil con un solo clic o atajo de teclado.")
bullet(doc, "Detectar automáticamente cuando conectas o desconectas un monitor y restaurar el perfil activo.")
bullet(doc, "Usar la app en español o inglés.")

h2(doc, "Requisitos del sistema")
add_table(doc,
    ["Requisito", "Detalle"],
    [
        ["Sistema operativo", "Windows 10 / Windows 11 (64-bit)"],
        ["Privilegios",       "Usuario estándar (ver nota sobre ventanas elevadas)"],
        ["Espacio en disco",  "~30 MB"],
        ["Dependencias",      "Ninguna — ejecutable autocontenido"],
    ],
    col_widths=[5, 10]
)

tip(doc, "Si necesitas mover ventanas de programas que se ejecutan como Administrador "
         "(Administrador de tareas, Omen Gaming Hub, etc.), ejecuta MonitorReminder "
         "también como Administrador.")

doc.add_page_break()

# ── 2. INSTALACIÓN ────────────────────────────────────────────────────────────
h1(doc, "2. Instalación")
h2(doc, "Descargar el ejecutable")
numbered(doc, "Ve a https://github.com/erickson558/monitorreminder/releases")
numbered(doc, "Descarga MonitorReminder.exe de la versión más reciente.")
numbered(doc, "Coloca el archivo en la carpeta que prefieras (no requiere instalador).")
numbered(doc, "Haz doble clic en MonitorReminder.exe para iniciarlo.")

tip(doc, "Puedes crear un acceso directo en el escritorio o añadirlo al inicio de Windows "
         "para que se ejecute automáticamente al encender tu computadora.")

h2(doc, "Primera ejecución")
body(doc,
    "Al ejecutar la app por primera vez se crea automáticamente el archivo config.json "
    "en la misma carpeta del ejecutable. Este archivo guarda todos tus perfiles y "
    "preferencias. No lo elimines a menos que quieras reiniciar la configuración."
)

doc.add_page_break()

# ── 3. INTERFAZ PRINCIPAL ─────────────────────────────────────────────────────
h1(doc, "3. Interfaz principal")
body(doc,
    "La ventana principal tiene cuatro zonas: el área Hero en la parte superior, "
    "el panel izquierdo de perfiles, el panel derecho de información y automatización, "
    "y la barra de estado en la parte inferior."
)

h2(doc, "3.1 Área Hero (parte superior)")
body(doc,
    "Contiene el nombre y subtítulo de la aplicación, el selector de idioma, el botón "
    "de soporte, los botones de monitoreo y salida, y el indicador de estado del sistema."
)
add_table(doc,
    ["Elemento", "Función"],
    [
        ["Selector de idioma",    "Cambia entre Español e Inglés. La app se reinicia para aplicar el cambio."],
        ["Cómprame una cerveza",  "Abre la página de donación vía PayPal para apoyar el proyecto."],
        ["Iniciar monitoreo",     "Activa el watcher de cambios de pantalla. El texto cambia a 'Detener monitoreo' mientras está activo."],
        ["Salir",                 "Guarda el estado y cierra la aplicación correctamente."],
        ["SISTEMA EN LINEA",      "Indicador visual de que la app está funcionando correctamente."],
    ],
    col_widths=[5, 10]
)

h2(doc, "3.2 Panel izquierdo — Perfiles guardados")
body(doc,
    "Muestra hasta 5 tarjetas de perfil. La tarjeta activa se resalta en verde. "
    "Cada tarjeta muestra el nombre del perfil y la fecha/hora del último guardado."
)
add_table(doc,
    ["Botón", "Atajo", "Función"],
    [
        ["Guardar perfil",   "Ctrl+S", "Captura las ventanas visibles y guarda sus posiciones en el perfil seleccionado."],
        ["Restaurar perfil", "Ctrl+R", "Mueve todas las ventanas guardadas a sus posiciones originales."],
        ["Nombrar perfil",   "—",      "Guarda el nombre escrito en el campo de texto como nombre del perfil activo."],
    ],
    col_widths=[4, 2.5, 9]
)

body(doc, "Campo de nombre de perfil:")
bullet(doc, "Escribe el nombre directamente en el campo de texto.")
bullet(doc, "Presiona Enter o haz clic fuera del campo para confirmar.")
bullet(doc, "Máximo 40 caracteres.")

h2(doc, "3.3 Panel derecho — Radar de pantallas y Automatización")

h3(doc, "Radar de pantallas")
body(doc,
    "Muestra en tiempo real el nombre, resolución y coordenadas de cada monitor conectado. "
    "Se actualiza automáticamente cuando conectas o desconectas una pantalla."
)

h3(doc, "¿Cómo se usa?")
body(doc, "Tarjeta de instrucciones rápidas integrada en la interfaz:")
numbered(doc, "Abre tus programas y acomódalos como quieras en tus monitores.")
numbered(doc, "Selecciona un perfil y presiona Guardar perfil.")
numbered(doc, "Cuando quieras volver al mismo orden, presiona Restaurar perfil.")
tip(doc, "Activa Iniciar monitoreo para restaurar automáticamente al cambiar o reconectar monitores.")

h3(doc, "Automatización")
add_table(doc,
    ["Opción", "Descripción"],
    [
        ["Auto iniciar proceso al abrir la aplicación",
         "Si está marcado, el watcher se activa automáticamente cada vez que abres MonitorReminder."],
        ["Cerrar la app automáticamente después de: N segundos",
         "Si está marcado, la app se cierra sola después del tiempo configurado. Útil si sólo la abres para restaurar un perfil."],
        ["Campo de segundos",
         "Número de segundos para el cierre automático. Rango válido: 5 a 3600 segundos."],
    ],
    col_widths=[6, 9]
)

h3(doc, "Telemetría de ventanas")
body(doc,
    "Lista de texto en formato monoespaciado que muestra las ventanas guardadas en el "
    "perfil activo (proceso | título de ventana). Muestra hasta 40 ventanas. "
    "Se actualiza automáticamente al guardar o cambiar de perfil."
)

h2(doc, "3.4 Barra de estado")
body(doc,
    "Barra horizontal en la parte inferior que muestra mensajes de retroalimentación "
    "en tiempo real, el contador de cierre automático (si está activo), y la versión "
    "de la aplicación."
)
add_table(doc,
    ["Mensaje", "Significado"],
    [
        ["Listo", "La app está en espera, sin operaciones activas."],
        ["Monitoreando cambios de pantalla", "El watcher está activo y vigilando conexiones/desconexiones de monitores."],
        ["Perfil restaurado (posición exacta): movidas X, ya correctas Y, no encontradas Z, fallidas W",
         "Resumen detallado del último restore. 'Posición exacta' indica que la configuración de monitores no cambió desde el guardado."],
        ["Perfil restaurado (posición proporcional): ...",
         "La configuración de monitores cambió; las ventanas se colocaron en posición proporcional relativa al nuevo tamaño."],
        ["✓ Todo está como debe en 'Perfil' (modo)",
         "Todas las ventanas ya están en su posición correcta; no se movió nada."],
        ["Auto-restaurado tras cambio de pantalla (...)",
         "El watcher detectó un cambio de monitor y restauró automáticamente."],
        ["Ocurrió un error. Revisa log.txt", "Se produjo un error interno. Consulta el archivo log.txt para detalles."],
    ],
    col_widths=[7, 8]
)

doc.add_page_break()

# ── 4. GESTIÓN DE PERFILES ────────────────────────────────────────────────────
h1(doc, "4. Gestión de perfiles")

h2(doc, "¿Qué es un perfil?")
body(doc,
    "Un perfil es una instantánea de todas las ventanas visibles en tu escritorio: "
    "qué programas estaban abiertos, en qué monitor, en qué posición y con qué tamaño. "
    "MonitorReminder guarda hasta 5 perfiles independientes que puedes usar para diferentes "
    "configuraciones de trabajo."
)
body(doc, "Cada perfil almacena:")
bullet(doc, "Nombre personalizado.")
bullet(doc, "Fecha y hora del último guardado.")
bullet(doc, "Por cada ventana: título, proceso, clase, monitor asignado, posición relativa, estado (normal/maximizado/minimizado).")
bullet(doc, "Firma de configuración de monitores (para detectar si cambiaron las pantallas).")

h2(doc, "Guardar un perfil")
numbered(doc, "Abre todos los programas que deseas recordar y acomódalos en tus monitores.")
numbered(doc, "Haz clic en la tarjeta del perfil donde quieres guardar.")
numbered(doc, "Presiona el botón Guardar perfil o usa Ctrl+S.")
numbered(doc, "Si el perfil ya tenía ventanas guardadas, aparecerá un diálogo de confirmación. Confirma para sobrescribir.")
numbered(doc, "La barra de estado mostrará cuántas ventanas se guardaron y restauraron.")

warn(doc, "Al guardar un perfil se restauran inmediatamente las posiciones para que todo quede alineado. "
          "Esto moverá tus ventanas activas.")

tip(doc, "Escribe un nombre descriptivo antes de guardar, por ejemplo: "
         "'Trabajo - 2 monitores' o 'Casa - Solo laptop'.")

h2(doc, "Restaurar un perfil")
numbered(doc, "Selecciona el perfil haciendo clic en su tarjeta.")
numbered(doc, "Presiona Restaurar perfil o usa Ctrl+R.")
numbered(doc, "MonitorReminder buscará cada ventana guardada y la moverá a su posición original.")
numbered(doc, "La barra de estado mostrará el resumen: ventanas movidas, ya correctas, no encontradas y fallidas.")

body(doc, "Modos de restauración:")
add_table(doc,
    ["Modo", "Cuándo se usa", "Comportamiento"],
    [
        ["Posición exacta",
         "La configuración de monitores no cambió desde el guardado.",
         "Cada ventana se coloca en exactamente los mismos píxeles guardados."],
        ["Posición proporcional",
         "Uno o más monitores cambiaron de resolución, nombre o posición.",
         "Las ventanas se recolocan proporcionalmente al nuevo tamaño del monitor."],
    ],
    col_widths=[3.5, 5, 6.5]
)

h2(doc, "Renombrar un perfil")
numbered(doc, "Selecciona el perfil que deseas renombrar.")
numbered(doc, "Escribe el nuevo nombre en el campo de texto en la parte inferior del panel izquierdo.")
numbered(doc, "Presiona Enter o haz clic fuera del campo.")
body(doc, "El nombre se guarda automáticamente. No se pide confirmación para renombrar.")

h2(doc, "Cambiar entre perfiles")
body(doc,
    "Haz clic en cualquier tarjeta de perfil para seleccionarla. La tarjeta activa se "
    "resalta en verde. El cambio de perfil no restaura las ventanas automáticamente — "
    "usa Restaurar perfil para aplicarlo."
)

doc.add_page_break()

# ── 5. MONITOR AUTOMÁTICO ─────────────────────────────────────────────────────
h1(doc, "5. Monitor automático (Watcher)")
body(doc,
    "El watcher es un proceso en segundo plano que vigila cambios en la configuración "
    "de pantallas. Cuando detecta que conectaste, desconectaste o reconfiguraste un "
    "monitor, espera 3 segundos y luego restaura automáticamente el perfil activo."
)

h2(doc, "Activar y desactivar")
bullet(doc, "Botón Iniciar monitoreo en el área Hero: activa el watcher manualmente.")
bullet(doc, "Botón Detener monitoreo: desactiva el watcher.")
bullet(doc, "Casilla 'Auto iniciar proceso al abrir la aplicación': si está marcada, el watcher inicia solo al abrir MonitorReminder.")

h2(doc, "Comportamiento al detectar un cambio de monitor")
numbered(doc, "El watcher detecta el evento de cambio de pantalla.")
numbered(doc, "Actualiza el Radar de pantallas en la interfaz.")
numbered(doc, "Espera 3 segundos para que Windows termine de reposicionar las ventanas.")
numbered(doc, "Restaura el perfil activo en segundo plano (la UI no se congela).")
numbered(doc, "La barra de estado muestra el resumen del auto-restore.")

tip(doc, "Si la app está ocupada con otro auto-restore cuando llega un nuevo evento de monitor, "
         "la barra de estado mostrará 'Sincronizando cambios de monitor, espera un momento...' "
         "y el segundo evento se ignorará para evitar conflictos.")

doc.add_page_break()

# ── 6. CIERRE AUTOMÁTICO ──────────────────────────────────────────────────────
h1(doc, "6. Cierre automático")
body(doc,
    "La función de cierre automático permite que MonitorReminder se cierre solo después "
    "de un número configurado de segundos. Es útil si usas la app solo para restaurar "
    "un perfil y no necesitas mantenerla abierta."
)

h2(doc, "Configuración")
numbered(doc, "Marca la casilla 'Cerrar la app automáticamente después de:' en el panel de Automatización.")
numbered(doc, "Escribe el número de segundos en el campo de texto (rango: 5 a 3600).")
numbered(doc, "La cuenta regresiva aparece en la barra de estado: N s / restantes s.")
numbered(doc, "Cuando llega a cero, la app se cierra guardando el estado.")

body(doc,
    "Para desactivarlo, desmarca la casilla. El contador se resetea cada vez que "
    "activas o cambias el valor de segundos."
)

doc.add_page_break()

# ── 7. ATAJOS DE TECLADO ──────────────────────────────────────────────────────
h1(doc, "7. Atajos de teclado")
add_table(doc,
    ["Atajo", "Acción"],
    [
        ["Ctrl + S", "Guardar perfil activo"],
        ["Ctrl + R", "Restaurar perfil activo"],
        ["Ctrl + Q", "Salir de la aplicación"],
        ["Enter (en campo de nombre)", "Confirmar nombre del perfil"],
    ],
    col_widths=[4, 11]
)
tip(doc, "Los atajos funcionan desde cualquier lugar de la ventana, incluso si el foco "
         "está en un campo de texto.")

doc.add_page_break()

# ── 8. CONFIGURACIÓN DE IDIOMA ────────────────────────────────────────────────
h1(doc, "8. Configuración de idioma")
body(doc,
    "MonitorReminder está completamente traducido al español y al inglés. Para cambiar "
    "el idioma usa el menú desplegable en la esquina superior derecha del área Hero."
)
add_table(doc,
    ["Opción", "Idioma"],
    [
        ["es", "Español (predeterminado)"],
        ["en", "English"],
    ],
    col_widths=[3, 12]
)
warn(doc, "Al cambiar el idioma la aplicación se reinicia automáticamente para refrescar "
          "todas las etiquetas. Las ventanas y perfiles no se ven afectados.")

doc.add_page_break()

# ── 9. COMPATIBILIDAD ─────────────────────────────────────────────────────────
h1(doc, "9. Compatibilidad de aplicaciones")
body(doc,
    "MonitorReminder puede restaurar la gran mayoría de ventanas de Windows. "
    "A continuación se detalla el comportamiento con categorías especiales de aplicaciones."
)

h2(doc, "9.1 Aplicaciones con títulos dinámicos — Totalmente soportadas desde V0.3.2")
body(doc,
    "Estas aplicaciones cambian su título de ventana con el archivo o pestaña activa. "
    "MonitorReminder usa una búsqueda por sufijo compartido para encontrarlas aunque el "
    "título haya cambiado."
)
add_table(doc,
    ["Aplicación", "Ejemplo de título guardado", "Cómo se encuentra"],
    [
        ["Visual Studio Code", "foo.py - mi-proyecto - Visual Studio Code", "Sufijo: 'Visual Studio Code'"],
        ["Microsoft Edge",     "GitHub - Microsoft Edge",                    "Sufijo: 'Microsoft Edge'"],
        ["Brave",              "YouTube - Brave",                            "Sufijo: 'Brave'"],
        ["Foxit PDF Reader",   "documento.pdf - Foxit PDF Reader",           "Sufijo: 'Foxit PDF Reader'"],
        ["Postman",            "Mi API - Postman",                           "Sufijo: 'Postman'"],
        ["Google Chrome",      "Inbox - Google Chrome",                      "Sufijo: 'Google Chrome'"],
    ],
    col_widths=[4, 7, 4]
)

h2(doc, "9.2 Aplicaciones UWP (Windows Store)")
body(doc,
    "Aplicaciones como Configuración de Windows usan una clase de ventana especial "
    "(`ApplicationFrameWindow`). MonitorReminder usa `SetWindowPlacement` para "
    "reposicionarlas, que es más compatible con este tipo de apps."
)
add_table(doc,
    ["Aplicación", "Comportamiento"],
    [
        ["Configuración de Windows", "Generalmente funciona. Si no se mueve, cierra y vuelve a abrir la app."],
        ["Microsoft Store",          "Puede resistir el reposicionamiento según el estado de la app."],
        ["Calculadora, Fotos, etc.", "Funcionan en la mayoría de los casos."],
    ],
    col_widths=[5, 10]
)

h2(doc, "9.3 Procesos elevados (requieren ejecutar como Administrador)")
body(doc,
    "Windows bloquea que un proceso normal mueva las ventanas de un proceso con "
    "privilegios de administrador (UIPI — User Interface Privilege Isolation). "
    "MonitorReminder detecta este bloqueo y lo reporta en el log."
)
add_table(doc,
    ["Aplicación", "Razón por la que requiere admin"],
    [
        ["Administrador de tareas (Taskmgr.exe)", "Siempre se ejecuta elevado"],
        ["Omen Gaming Hub",                        "Se ejecuta con privilegios elevados"],
        ["Hard Disk Sentinel",                     "Requiere acceso a hardware de discos"],
        ["Cualquier app iniciada con 'Ejecutar como administrador'", "Por definición, proceso elevado"],
    ],
    col_widths=[6, 9]
)
warn(doc, "Para mover estas ventanas: clic derecho en MonitorReminder.exe → "
          "Propiedades → Compatibilidad → Marcar 'Ejecutar este programa como administrador'.")

tip(doc, "Después de restaurar, la barra de estado muestra las ventanas 'fallidas'. "
         "Si una ventana importante siempre falla, probablemente necesitas ejecutar como admin.")

doc.add_page_break()

# ── 10. SOLUCIÓN DE PROBLEMAS ─────────────────────────────────────────────────
h1(doc, "10. Solución de problemas")

add_table(doc,
    ["Problema", "Causa probable", "Solución"],
    [
        ["Una ventana no regresa a su posición",
         "La app fue encontrada pero no pudo moverse (proceso elevado)",
         "Ejecuta MonitorReminder como Administrador."],
        ["Una ventana aparece como 'no encontrada'",
         "La aplicación estaba cerrada cuando se ejecutó el restore",
         "Abre la aplicación antes de restaurar."],
        ["VSCode / Edge / Brave no se mueven",
         "El título cambió; no se encontró coincidencia (versión anterior a 0.3.2)",
         "Actualiza a V0.3.2 o superior."],
        ["El monitoreo no detecta la conexión del monitor",
         "Algunos adaptadores de video no envían el evento de cambio de pantalla a tiempo",
         "Usa el botón Restaurar perfil manualmente después de conectar el monitor."],
        ["La ventana se restaura en la posición incorrecta",
         "El perfil fue guardado con una configuración de monitores diferente",
         "Reconnecta los monitores en la misma configuración y vuelve a guardar el perfil."],
        ["Error al iniciar: 'No se puede abrir config.json'",
         "El archivo de configuración está corrupto o bloqueado",
         "Elimina config.json y reinicia la app (los perfiles se pierden)."],
        ["La ventana de MonitorReminder no aparece",
         "La app puede estar minimizada o fuera de pantalla",
         "Busca el ícono en la barra de tareas o usa Alt+Tab."],
    ],
    col_widths=[4.5, 4.5, 6]
)

h2(doc, "Revisar el archivo log.txt")
body(doc,
    "MonitorReminder escribe registros detallados de todas sus operaciones en log.txt, "
    "ubicado en la misma carpeta que el ejecutable. El archivo incluye:"
)
bullet(doc, "Fecha y hora de cada operación.")
bullet(doc, "Número de ventanas capturadas/restauradas.")
bullet(doc, "Advertencias sobre ventanas elevadas (UIPI).")
bullet(doc, "Errores con el stack trace completo.")
tip(doc, "Abre log.txt con el Bloc de notas o cualquier editor de texto. "
         "Las entradas más recientes están al final del archivo.")

doc.add_page_break()

# ── 11. PREGUNTAS FRECUENTES ──────────────────────────────────────────────────
h1(doc, "11. Preguntas frecuentes (FAQ)")

add_table(doc,
    ["Pregunta", "Respuesta"],
    [
        ["¿Se pueden tener más de 5 perfiles?",
         "No. El límite actual es 5 perfiles. Puedes renombrarlos para diferentes configuraciones."],
        ["¿MonitorReminder captura las ventanas minimizadas?",
         "Sí. Las ventanas minimizadas se capturan con su posición normal (la que tendrían al restaurarse) y se restauran minimizadas."],
        ["¿Qué pasa si tengo dos ventanas de VSCode abiertas?",
         "MonitorReminder elige la que mejor coincide por sufijo de título. Si el sufijo es idéntico, toma la primera encontrada."],
        ["¿Funciona con múltiples escritorios virtuales de Windows?",
         "No. MonitorReminder trabaja con el escritorio activo. Las ventanas en otros escritorios virtuales no se capturan."],
        ["¿Se envían mis datos a algún servidor?",
         "No. MonitorReminder es 100% local. config.json y log.txt permanecen en tu computadora."],
        ["¿Puedo usar el perfil en otra computadora?",
         "El archivo config.json usa posiciones absolutas de píxeles. Funciona correctamente solo en la misma computadora con la misma configuración de monitores."],
        ["¿Por qué la app espera 3 segundos antes de auto-restaurar?",
         "Windows tarda unos segundos en reposicionar las ventanas tras un cambio de monitor. Esperar evita conflictos entre el restore de MonitorReminder y el de Windows."],
    ],
    col_widths=[6.5, 8.5]
)

doc.add_page_break()

# ── 12. ACERCA DE ─────────────────────────────────────────────────────────────
h1(doc, "12. Acerca de")
add_table(doc,
    ["Campo", "Valor"],
    [
        ["Nombre",            "MonitorReminder"],
        ["Versión",           "V0.3.2"],
        ["Autor",             "Synyster Rick"],
        ["Licencia",          "Apache 2.0"],
        ["Repositorio",       "https://github.com/erickson558/monitorreminder"],
        ["Releases",          "https://github.com/erickson558/monitorreminder/releases"],
        ["Año",               "2026"],
    ],
    col_widths=[4, 11]
)

doc.add_paragraph()
body(doc,
    "Si MonitorReminder te ha ahorrado tiempo y dolores de cabeza configurando tus "
    "monitores, considera apoyar el proyecto con una pequeña donación. Cada contribución "
    "ayuda a mantener actualizaciones y nuevas funciones.",
    color=TEAL
)
body(doc, "Donación vía PayPal: https://www.paypal.com/donate/?hosted_button_id=ZABFRXC2P3JQN",
     color=CYAN)

divider(doc)

footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f = footer_p.add_run("MonitorReminder V0.3.2  ·  Manual de Usuario  ·  2026")
r_f.font.size = Pt(9)
r_f.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
r_f.italic = True

# ── SAVE ─────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Manual generado: {OUT}")
